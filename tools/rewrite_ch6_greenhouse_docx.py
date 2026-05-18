from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path(r"D:\pychram\RAG2\output\target_greenhouse_report.docx")
OUT = Path(r"D:\pychram\RAG2\output\target_greenhouse_report_ch6_revised.docx")


PRODUCTS = [
    {
        "title": "6.1.1 轻量化环境智能控制终端（优先级：高）",
        "pain": "解决1-5亩小型温室中环境调控高频、依赖人工、响应不及时的问题。当前多数大棚仍停留在电动化或半自动阶段，通风、卷帘、保温、灌溉等操作需要农户多次到棚内巡查，极端天气下延迟调控会直接影响作物存活率、产量和品质稳定性。",
        "tech": "采用“传感器+无线网关+执行器适配模块+手机端控制”的轻量化改造路径，接入温湿度、光照、CO2、土壤湿度等传感器，并兼容卷帘机、风机、遮阳帘、水泵和电磁阀。边缘端负责阈值预警与弱网缓存，云端根据作物模型生成调控建议，先做自动提醒和一键执行，再逐步升级为闭环控制。",
        "market": "我国温室大棚面积大、存量改造空间明显，但真正实现智能化管理的比例仍偏低。大型智能温室系统价格高、施工重，难以覆盖中小农户；地方设施农业改造和农机补贴政策则为低成本、可快速部署的环境控制终端提供了推广窗口。",
        "rows": [
            ["成本", "单棚硬件成本", "2000-5000元/棚", "智能温室投资成本较五年前下降约35%，智能控制系统占智能温室投资的20%-30%"],
            ["部署", "安装与改造时间", "半天内完成，尽量不改棚体结构", "现有大棚多已有卷帘机、风机、水泵等电动设备，适合做增量控制改造"],
            ["控制效果", "环境超限响应时间", "5分钟内预警或执行", "设施蔬菜环境调控机械化率约47%-61%，仍存在明显人工干预空间"],
            ["ROI", "投资回收期", "1-2个种植季", "通过减少巡棚、降低极端天气损失、稳定品质来体现收益"],
        ],
        "commercial": "优先采用“硬件低门槛销售+年度运维服务+补贴申报协助”的模式，在设施蔬菜主产区选择合作社或示范户打样；成熟后拆分为基础版、增强版和多棚组网版，通过农资店、农机经销商和地方示范项目下沉推广。",
    },
    {
        "title": "6.1.2 便携式作物巡检与病虫害识别设备（优先级：高）",
        "pain": "解决中小农户病虫害识别依赖经验、巡检覆盖不全、错过最佳防治窗口的问题。设施环境湿度高、病虫害传播快，误判或晚判会带来减产和用药浪费；农户需要的是能直接给出风险等级、处理建议和用药注意事项的辅助决策工具。",
        "tech": "以手机APP或小程序为核心，结合AI视觉识别、病虫害知识库、作物生长阶段判断和本地农技规则。低配方案只依赖手机拍照，高配方案可增加微距镜头、诱虫板识别、固定摄像头或虫情测报灯，形成从“拍照诊断”到“区域预警”的阶梯式产品。",
        "market": "AI病虫害识别技术成熟度较高，识别准确率已有多个产品达到90%以上，软件化路径能显著降低初始采购门槛。与农资销售、农技服务和保险理赔结合后，产品不仅能解决诊断问题，还能形成用药、复诊、复购和数据沉淀闭环。",
        "rows": [
            ["识别能力", "田间实测准确率", "不低于90%", "虫情测报、农技识别类产品常见宣传准确率为90%+，部分产品达到94%-95%"],
            ["响应速度", "拍照到出结果时间", "5秒内", "移动端识别工具通常以秒级反馈作为用户体验基准"],
            ["覆盖范围", "作物/病虫害覆盖数量", "覆盖本地高频作物和100种以上常见病虫害", "成熟病虫害知识库可覆盖数百至上千类作物问题"],
            ["成本", "农户使用费用", "基础版0-200元/季，硬件增强版3000元以内", "纯软件和小程序可免费或低价获客，物联网测报灯等硬件单价约数千元"],
        ],
        "commercial": "短期用免费诊断或低价订阅获取样本和农户信任，中期与农资店、植保服务队、合作社绑定，按季收取AI诊断服务费或从农资成交中分成；大型基地可销售固定摄像头和虫情监测硬件，形成SaaS年费。",
    },
    {
        "title": "6.1.3 简易水肥一体化控制器（优先级：中）",
        "pain": "解决中小温室水肥管理粗放、肥料利用率低、人工配肥灌溉费时的问题。传统“大水大肥”既增加肥料和水电成本，也会造成作物品质波动和面源污染；但现有高配水肥系统价格高、安装复杂，超出大量中小农户承受能力。",
        "tech": "采用单通道或少通道控制器、EC/pH传感器、流量计、电磁阀、施肥泵和简化作物配方库，先实现定时定量灌溉、肥液浓度监测和异常提醒。云端根据作物阶段、天气和土壤湿度生成水肥建议，大型基地版本再扩展到多区控制和历史数据分析。",
        "market": "水肥一体化应用面积已突破1.5亿亩，但中小设施温室渗透率仍低。节水、节肥、稳产的价值容易量化，且多地节水农业、设施农业改造和农机补贴可为推广提供政策支撑，因此适合作为第二梯队规模化产品。",
        "rows": [
            ["成本", "1-5亩单套设备成本", "1000-3000元", "简易设备可做到亩均约200元，单通道机型成本约为多通道方案的50%-60%"],
            ["节水", "较传统漫灌节水比例", "不低于30%", "水肥一体化常见节水效果为30%-50%"],
            ["节肥", "较传统施肥节肥比例", "不低于20%", "水肥利用率可由30%-40%提升至70%以上，部分作物化肥成本下降约25%"],
            ["增产", "较传统模式增产比例", "不低于10%", "番茄、菜豆等设施作物案例中增产可达12%以上"],
        ],
        "commercial": "以“控制器硬件+水肥配方服务+农资联动”为主，基础版面向中小农户一次性低价销售，增强版收取年度算法和远程运维费。可与肥料企业、灌溉设备商和合作社联合推广，用节水节肥数据支撑复购。",
    },
    {
        "title": "6.1.4 辅助采摘/搬运设备（优先级：中）",
        "pain": "解决采收和棚内搬运环节人工成本高、劳动强度大、旺季用工不稳定的问题。设施蔬菜采收机械化率仍处于低位，果菜和瓜菜对人工依赖强；全自主采摘机器人短期内成本和适应性压力较大，中小农户更需要半自动、省力化的辅助设备。",
        "tech": "优先发展电动轨道运输车、跟随式搬运车、升降采摘平台和轻量化周转箱系统，而不是直接切入高成本全自主采摘机器人。技术重点放在窄棚通行、稳定载重、低速避障、可拆装轨道、换电续航和人机协同流程。",
        "market": "采收人工成本通常占生产成本的30%-60%，草莓、番茄等高价值设施作物痛点更强。半自动搬运设备技术成熟、价格低、ROI更容易算清，适合先从大型基地和标准化棚型切入，再通过租赁或共享服务覆盖小户。",
        "rows": [
            ["成本", "单套设备投资", "轨道车3000-8000元；智能搬运车2-5万元", "已有轨道运输、智能搬运车案例显示可在园区和温室场景落地"],
            ["效率", "搬运效率提升", "不低于3倍", "部分智能搬运设备可实现3-4倍效率提升并支持连续作业"],
            ["载重", "单次最大载重", "不低于200公斤", "棚内运输需覆盖果筐、肥料、农资和采收周转箱等常见负载"],
            ["降本", "采收/搬运人工成本下降", "下降30%-50%", "采收人工成本当前可占30%-60%，省工空间明确"],
        ],
        "commercial": "不宜先做重资产销售，可由合作社、村集体、农业服务公司购买后按棚、按天或按季出租。对大型基地销售整套物流动线方案，对中小农户提供共享设备和旺季托管服务。",
    },
    {
        "title": "6.1.5 智能授粉辅助设备（优先级：探索方向）",
        "pain": "解决设施经济作物授粉窗口短、人工授粉质量不稳定、熊蜂授粉受环境和农药影响的问题。授粉直接影响坐果率、果形和商品率，但该环节对识别精度、作业时机和环境适应要求高，当前更适合技术跟踪和定向试点。",
        "tech": "采用AI视觉识别花朵状态，结合微型机械臂、气流或振动装置完成非接触授粉，并通过路线规划和花期数据记录形成闭环。短期应聚焦番茄、草莓、甜瓜等棚型标准、花朵特征清晰、单果价值较高的作物。",
        "market": "授粉机器人已有示范案例，但成本、通用性、稳定性和田间维护仍未达到中小农户普及条件。该方向适合3-5年内面向中等以上基地或育种企业试点，待成本下降和作物模型成熟后再进入普惠市场。",
        "rows": [
            ["成功率", "田间授粉成功率", "商业化目标不低于95%", "现有AI授粉方案公开案例约82%-98%，稳定性仍需持续验证"],
            ["效率", "单花作业时间", "5秒/朵以内", "部分样机单花作业时间约10-15秒，仍需提升并行作业能力"],
            ["成本", "单台设备目标成本", "中期降至3-5万元", "当前更适合大型温室或科研育种场景，普惠价格带尚未形成"],
            ["适用性", "首批适配作物", "番茄、草莓、甜瓜等标准化设施作物", "这些作物价值高、棚型相对稳定，更适合早期验证"],
        ],
        "commercial": "短期不建议面向中小农户直接销售，可采用“试点项目+按季服务+科研/育种合作”的商业化路径。先与大型基地、种业企业和地方示范项目共建样板，验证坐果率、商品率和人工替代数据。",
    },
]


BLIND_SPOTS = [
    (
        "产品盲点：现有供给的结构性缺失",
        "市场并不缺“智慧农业概念产品”，缺的是能进入中小温室真实经营边界的产品：低成本、轻安装、低培训、可单棚部署、可在1-2个种植季内证明收益。当前供给一端是大型农场系统级方案，另一端是功能单薄的传感器或单机设备，中间缺少围绕环境控制、巡检、水肥、搬运等高频环节设计的标准化轻量套件。",
    ),
    (
        "商业模式盲点：服务方式的缺失",
        "一次性硬件销售与中小农户现金流、风险承受力和技术信任形成错配。真正需要补齐的是租赁、共享、按季订阅、按次巡检、农资联动、补贴申报协助和农忙期快速运维等服务方式，把设备从“固定资产采购”转化为“可试用、可分期、可复购的生产服务”。",
    ),
    (
        "数据/算法盲点：技术落地的基础瓶颈",
        "设施温室存在强区域性和小气候差异，同一作物在不同棚型、季节和管理习惯下的数据分布差异明显。病虫害识别、长势判断、水肥策略和授粉窗口都需要本地化样本和持续标注；如果没有作物全周期图像、环境、作业和产量数据闭环，AI模型容易停留在演示效果，难以支撑稳定交付和复购。",
    ),
    (
        "盲点间的逻辑关联与破局建议",
        "三类盲点是递进关系：产品不够轻，用户不愿试；付费方式不够灵活，用户不敢买；数据闭环不足，产品用不好、复购弱。破局路径应从单点高频场景切入，先在标准化作物和设施蔬菜主产区建立示范棚，用成本、节工、增产、减损和回收期作为硬指标；再通过合作社、农资店、农机经销商和地方示范项目扩散，持续沉淀本地作物数据，形成“产品套件+服务网络+算法数据”的复合壁垒。",
    ),
]


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def set_paragraph_text(paragraph, text: str, style: str | None = None) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    paragraph.add_run(text)


def paragraph_text_from_element(element) -> str:
    return "".join(t.text for t in element.iter(qn("w:t")) if t.text).strip()


def insert_para_before(ref_para, text: str = "", style: str = "Normal", *, bold_prefix: str = ""):
    p = ref_para.insert_paragraph_before(style=style)
    if style == "TBBT":
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(3)
    if not bold_prefix:
        p.add_run(text)
        return p
    run = p.add_run(bold_prefix)
    run.bold = True
    p.add_run(text)
    return p


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Cm(width)


def insert_table_before(doc, ref_para, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Table Grid"
    headers = ["维度", "关键指标", "目标值", "市场参考数据"]
    set_table_widths(table, [2.2, 3.0, 3.0, 6.2])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, font_size=8.8)
        shade_cell(cell, "D9EAF7")
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), text, font_size=8.3)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    ref_para._p.addprevious(table._tbl)


def update_toc(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("第六章 AI辅助产品方向") or text.startswith("第六章 产品方向分析"):
            if p.style.name.startswith("toc"):
                set_paragraph_text(p, "第六章 产品方向分析\t42", p.style.name)
        elif text.startswith("6.1 重点产品方向"):
            if p.style.name.startswith("toc"):
                set_paragraph_text(p, "6.1 重点产品方向\t42", p.style.name)
        elif text.startswith("6.2 市场盲点"):
            if p.style.name.startswith("toc"):
                set_paragraph_text(p, "6.2 市场盲点梳理\t51", p.style.name)
        elif text.startswith("第七章 结论与投资建议"):
            if p.style.name.startswith("toc"):
                set_paragraph_text(p, "第七章 结论与投资建议\t53", p.style.name)
        elif text.startswith("7.1 核心结论"):
            if p.style.name.startswith("toc"):
                set_paragraph_text(p, "7.1 核心结论\t53", p.style.name)
        elif text.startswith("7.2 投资建议"):
            if p.style.name.startswith("toc"):
                set_paragraph_text(p, "7.2 投资建议\t53", p.style.name)


def rewrite_chapter(doc: Document) -> None:
    chapter_para = None
    next_chapter_para = None
    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style is not None else ""
        if style_name == "Heading 1" and text.startswith("第六章"):
            chapter_para = p
        elif style_name == "Heading 1" and text.startswith("第七章"):
            next_chapter_para = p
            break
    if chapter_para is None or next_chapter_para is None:
        raise RuntimeError("Could not locate chapter boundaries.")

    set_paragraph_text(chapter_para, "第六章 产品方向分析", "Heading 1")

    body = doc._body._element
    start_idx = list(body).index(chapter_para._p)
    end_idx = list(body).index(next_chapter_para._p)
    for element in list(body)[start_idx + 1 : end_idx]:
        body.remove(element)

    ref = next_chapter_para
    insert_para_before(ref, "6.1 重点产品方向", "Heading 2")
    insert_para_before(
        ref,
        "本节按“先高频刚需、再效率改善、最后技术探索”的优先级排序。优先级判断主要依据四项条件：痛点是否高频且可量化、技术是否已经可工程化、补贴后成本是否进入目标用户承受区间、商业化路径是否能形成复购或服务收入。",
    )
    for idx, product in enumerate(PRODUCTS, start=20):
        insert_para_before(ref, product["title"], "Heading 3")
        insert_para_before(ref, product["pain"], bold_prefix="痛点对应：")
        insert_para_before(ref, product["tech"], bold_prefix="技术路径：")
        insert_para_before(ref, product["market"], bold_prefix="市场机会：")
        insert_para_before(ref, f"图表 {idx}：{product['title'].split(' ', 1)[1].split('（', 1)[0]}产品设计关键指标", "TBBT")
        insert_table_before(doc, ref, product["rows"])
        insert_para_before(ref, product["commercial"], bold_prefix="商业化路径：")

    insert_para_before(ref, "6.2 市场盲点梳理", "Heading 2")
    for title, body_text in BLIND_SPOTS:
        insert_para_before(ref, title, "Heading 3")
        insert_para_before(ref, body_text)


def main() -> None:
    doc = Document(SRC)
    update_toc(doc)
    rewrite_chapter(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
